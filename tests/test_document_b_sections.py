"""Tests for version-bound Document B section persistence."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from job_application_copilot.config import AppSettings
from job_application_copilot.domain import ReferenceAssetType
from job_application_copilot.repositories import (
    DocumentBSectionNotFoundError,
    create_database,
)
from job_application_copilot.services import (
    DocumentBSectionError,
    DocumentBSectionService,
    ReferenceAssetStorageService,
)
from job_application_copilot.services.database_bootstrap import initialize_database


def make_docx(*, heading: str = "CV generation rules", text: str = "Use evidence.") -> bytes:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_hierarchical_docx() -> bytes:
    document = Document()
    document.add_heading("Professional summary library", level=1)
    document.add_heading("Applied AI", level=2)
    document.add_heading("Core claim", level=3)
    document.add_paragraph("Applied AI claim.")
    document.add_heading("When to use", level=3)
    document.add_paragraph("Use for deployment roles.")
    document.add_heading("Head of Data", level=2)
    document.add_heading("Core claim", level=3)
    document.add_paragraph("Head of Data claim.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_context(
    tmp_path: Path,
) -> tuple[
    AppSettings,
    ReferenceAssetStorageService,
    DocumentBSectionService,
]:
    settings = AppSettings(_env_file=None, data_dir=tmp_path / "data")
    settings.database_path.parent.mkdir(parents=True)
    initialize_database(settings.database_path)
    database = create_database(settings.database_path)
    return (
        settings,
        ReferenceAssetStorageService(database, settings),
        DocumentBSectionService(database, settings),
    )


def test_extracts_stores_and_retrieves_exact_document_b_version(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    extracted = service.extract_and_store(version.version)
    listed = service.list_sections(version.version)
    required = service.require_section(version.version, "cv-generation-rules")

    assert extracted == listed
    assert required.document_b_version == version.version
    assert required.heading_title == "CV generation rules"
    assert required.section_text == "Use evidence."


def test_reextraction_replaces_sections_without_duplicates(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    first = service.extract_and_store(version.version)
    second = service.extract_and_store(version.version)

    assert first == second
    assert len(service.list_sections(version.version)) == 1


def test_first_read_extracts_legacy_version_without_section_rows(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    sections = service.list_sections(version.version)

    assert [section.section_id for section in sections] == ["cv-generation-rules"]
    assert service.list_sections(version.version) == sections


def test_versions_can_reuse_same_stable_section_id(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    first = storage.store(
        filename="document-b-1.docx",
        content=make_docx(text="First version."),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )
    second = storage.store(
        filename="document-b-2.docx",
        content=make_docx(text="Second version."),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    service.extract_and_store(first.version)
    service.extract_and_store(second.version)

    assert (
        service.require_section(first.version, "cv-generation-rules").section_text
        == "First version."
    )
    assert (
        service.require_section(second.version, "cv-generation-rules").section_text
        == "Second version."
    )


def test_rejects_tampered_retained_document(tmp_path: Path) -> None:
    settings, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )
    stored_path = settings.reference_folder / version.file_path
    stored_path.write_bytes(make_docx(text="Tampered."))

    with pytest.raises(DocumentBSectionError, match="recorded hash"):
        service.extract_and_store(version.version)


def test_missing_section_has_version_specific_error(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )
    service.extract_and_store(version.version)

    with pytest.raises(DocumentBSectionNotFoundError, match="missing"):
        service.require_section(version.version, "missing")


def test_lists_selected_heading_and_descendants_until_next_peer(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_hierarchical_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    tree = service.list_section_tree(
        version.version,
        "professional-summary-library-applied-ai",
    )

    assert [section.section_id for section in tree] == [
        "professional-summary-library-applied-ai",
        "professional-summary-library-applied-ai-core-claim",
        "professional-summary-library-applied-ai-when-to-use",
    ]
    assert "Head of Data claim." not in "\n".join(section.section_text for section in tree)


def test_leaf_section_tree_contains_only_selected_section(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_hierarchical_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    tree = service.list_section_tree(
        version.version,
        "professional-summary-library-applied-ai-core-claim",
    )

    assert len(tree) == 1
    assert tree[0].section_text == "Applied AI claim."


def test_missing_section_tree_has_version_specific_error(tmp_path: Path) -> None:
    _, storage, service = make_context(tmp_path)
    version = storage.store(
        filename="document-b.docx",
        content=make_docx(),
        asset_key="document-b",
        asset_type=ReferenceAssetType.DOCUMENT,
        name="Document B",
    )

    with pytest.raises(DocumentBSectionNotFoundError, match="missing"):
        service.list_section_tree(version.version, "missing")
