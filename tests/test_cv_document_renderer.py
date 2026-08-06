"""Tests for deterministic English DOCX rendering and local output."""

from datetime import date
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from job_application_copilot.config import AppSettings
from job_application_copilot.documents import (
    CvDocumentRenderError,
    render_cv_template,
    validate_docx,
)
from job_application_copilot.domain import (
    CvTemplateManifest,
    CvTemplateManifestStatus,
    CvTemplateSlotKind,
    CvTemplateSlotMapping,
    FinalCvOutput,
)
from job_application_copilot.repositories import Database, create_database
from job_application_copilot.services import CvDocumentRendererService, CvTemplateManifestService
from job_application_copilot.services.database_bootstrap import initialize_database


def final_output() -> FinalCvOutput:
    return FinalCvOutput.model_validate(
        {
            "opening_title": {"placeholder": "[OPENING_TITLE]", "content": "Architecture leader"},
            "opening_profile": {"placeholder": "[OPENING_PROFILE]", "content": "Evidence-led."},
            "experience": [
                {
                    "placeholder": "[EXPERIENCE_CURRENT]",
                    "title": {"placeholder": "[CURRENT_TITLE]", "content": "Technical Lead"},
                    "introduction": "Led a multidisciplinary team.",
                    "bullets": ["Designed durable systems.", "Improved delivery practices."],
                }
            ],
            "skills": {
                "placeholder": "[SKILLS]",
                "entries": [
                    {"name": "Architecture", "content": "APIs and integration."},
                    {"name": "Leadership", "content": "Team delivery."},
                ],
            },
        }
    )


def manifest() -> CvTemplateManifest:
    return CvTemplateManifest(
        template_asset_id=1,
        status=CvTemplateManifestStatus.CONFIRMED,
        placeholders=(
            "[OPENING_TITLE]",
            "[OPENING_PROFILE]",
            "[CURRENT_TITLE]",
            "[EXPERIENCE_CURRENT]",
            "[SKILLS]",
        ),
        slots=(
            CvTemplateSlotMapping(
                placeholder="[OPENING_TITLE]", kind=CvTemplateSlotKind.OPENING_TITLE
            ),
            CvTemplateSlotMapping(
                placeholder="[OPENING_PROFILE]", kind=CvTemplateSlotKind.OPENING_PROFILE
            ),
            CvTemplateSlotMapping(
                placeholder="[CURRENT_TITLE]",
                kind=CvTemplateSlotKind.EXPERIENCE_TITLE,
                experience_target="Current",
            ),
            CvTemplateSlotMapping(
                placeholder="[EXPERIENCE_CURRENT]",
                kind=CvTemplateSlotKind.EXPERIENCE,
                experience_target="Current",
            ),
            CvTemplateSlotMapping(placeholder="[SKILLS]", kind=CvTemplateSlotKind.SKILLS),
        ),
    )


def template_bytes(*, inline_title: bool = False) -> bytes:
    document = Document()
    static = document.add_paragraph("Daniel Brule | London")
    static.runs[0].bold = True
    title = document.add_paragraph("Role: [OPENING_TITLE]" if inline_title else "[OPENING_TITLE]")
    title.runs[0].italic = True
    document.add_paragraph("[OPENING_PROFILE]")
    document.add_paragraph("[CURRENT_TITLE]")
    bullets = document.add_paragraph("[EXPERIENCE_CURRENT]")
    bullets.style = "List Bullet"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "[SKILLS]"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_renders_generated_slots_without_changing_template_owned_content() -> None:
    rendered = render_cv_template(template_bytes(), manifest=manifest(), output=final_output())

    validate_docx("rendered.docx", rendered)
    document = Document(BytesIO(rendered))

    assert document.paragraphs[0].text == "Daniel Brule | London"
    assert document.paragraphs[0].runs[0].bold is True
    assert document.paragraphs[1].text == "Architecture leader"
    assert document.paragraphs[1].runs[0].italic is True
    assert [paragraph.text for paragraph in document.paragraphs[2:7]] == [
        "Evidence-led.",
        "Technical Lead",
        "Led a multidisciplinary team.",
        "Designed durable systems.",
        "Improved delivery practices.",
    ]
    assert [paragraph.style.name for paragraph in document.paragraphs[4:7]] == [
        "List Bullet",
        "List Bullet",
        "List Bullet",
    ]
    assert [paragraph.text for paragraph in document.tables[0].cell(0, 0).paragraphs] == [
        "Architecture: APIs and integration.",
        "Leadership: Team delivery.",
    ]


def test_rejects_generated_placeholder_embedded_in_static_template_text() -> None:
    with pytest.raises(CvDocumentRenderError, match="without surrounding text"):
        render_cv_template(
            template_bytes(inline_title=True), manifest=manifest(), output=final_output()
        )


def test_rejects_template_missing_a_configured_generated_placeholder() -> None:
    document = Document(BytesIO(template_bytes()))
    document.paragraphs[2]._element.getparent().remove(document.paragraphs[2]._element)
    buffer = BytesIO()
    document.save(buffer)

    with pytest.raises(CvDocumentRenderError, match="missing configured"):
        render_cv_template(buffer.getvalue(), manifest=manifest(), output=final_output())


@pytest.fixture
def renderer_service(tmp_path: Path) -> tuple[CvDocumentRendererService, CvTemplateManifestService]:
    settings = AppSettings(_env_file=None, data_dir=tmp_path / "data")
    settings.database_path.parent.mkdir(parents=True)
    initialize_database(settings.database_path)
    database: Database = create_database(settings.database_path)
    try:
        manifest_service = CvTemplateManifestService(database, settings)
        yield CvDocumentRendererService(database, settings), manifest_service
    finally:
        database.dispose()


def test_saves_non_overwriting_local_docx(
    renderer_service: tuple[CvDocumentRendererService, CvTemplateManifestService],
) -> None:
    service, manifest_service = renderer_service
    candidate = manifest_service.upload(filename="template.docx", content=template_bytes())
    manifest_service.confirm(template_asset_id=candidate.template_asset_id, slots=manifest().slots)

    first = service.render(final_output(), company="A/CME", generation_date=date(2026, 8, 6))
    second = service.render(final_output(), company="A/CME", generation_date=date(2026, 8, 6))

    assert first.name == "resume - Daniel Brule - 2026-08-06 - A-CME.docx"
    assert second.name == "resume - Daniel Brule - 2026-08-06 - A-CME (2).docx"
    validate_docx(first.name, first.read_bytes())
    validate_docx(second.name, second.read_bytes())
