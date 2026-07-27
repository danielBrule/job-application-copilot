"""Tests for deterministic local Document B heading extraction."""

from io import BytesIO

import pytest
from docx import Document

from job_application_copilot.documents import (
    DocumentBExtractionError,
    extract_document_b_sections,
)


def document_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extracts_preamble_hierarchy_tables_and_non_overlapping_text() -> None:
    document = Document()
    document.add_paragraph("CV Generation Guide", style="Title")
    document.add_paragraph("Purpose and version information.")
    preamble_table = document.add_table(rows=2, cols=2)
    preamble_table.cell(0, 0).text = "Version"
    preamble_table.cell(0, 1).text = "Status"
    preamble_table.cell(1, 0).text = "2026-06"
    preamble_table.cell(1, 1).text = "Working"
    document.add_heading("Professional summary library", level=1)
    document.add_paragraph("Library rules.")
    document.add_heading("Applied AI", level=2)
    document.add_heading("Purpose", level=3)
    document.add_paragraph("Applied AI purpose.")
    detail_table = document.add_table(rows=2, cols=2)
    detail_table.cell(0, 0).text = "Use"
    detail_table.cell(0, 1).text = "Avoid"
    detail_table.cell(1, 0).text = "Deployment roles"
    detail_table.cell(1, 1).text = "Unrelated roles"
    document.add_heading("Head of Data", level=2)
    document.add_heading("Purpose", level=3)
    document.add_paragraph("Head of Data purpose.")

    sections = extract_document_b_sections(document_bytes(document))

    assert [
        (
            section.section_id,
            section.heading_level,
            section.heading_title,
            section.sequence,
        )
        for section in sections
    ] == [
        ("document-preamble", 0, "Document preamble", 1),
        ("professional-summary-library", 1, "Professional summary library", 2),
        ("professional-summary-library-applied-ai", 2, "Applied AI", 3),
        ("professional-summary-library-applied-ai-purpose", 3, "Purpose", 4),
        ("professional-summary-library-head-of-data", 2, "Head of Data", 5),
        ("professional-summary-library-head-of-data-purpose", 3, "Purpose", 6),
    ]
    assert "Version | Status" in sections[0].section_text
    assert "2026-06 | Working" in sections[0].section_text
    assert sections[1].section_text == "Library rules."
    assert "Use | Avoid" in sections[3].section_text
    assert "Head of Data purpose." not in sections[3].section_text


def test_numbered_fallback_requires_heading_formatting() -> None:
    document = Document()
    document.add_heading("Rules", level=1)
    document.add_paragraph("1. Can the reader identify the role?")
    fallback = document.add_paragraph()
    fallback.add_run("2.3 Manual subsection").bold = True
    document.add_paragraph("Fallback content.")

    sections = extract_document_b_sections(document_bytes(document))

    assert [section.section_id for section in sections] == [
        "rules",
        "section-2-3",
    ]
    assert sections[0].section_text == "1. Can the reader identify the role?"
    assert sections[1].heading_number == "2.3"
    assert sections[1].heading_title == "Manual subsection"
    assert sections[1].heading_level == 2


def test_ignores_empty_heading_and_keeps_following_content_in_current_section() -> None:
    document = Document()
    document.add_heading("Bullet library", level=1)
    document.add_heading("", level=2)
    document.add_paragraph("Content after accidental empty heading.")

    sections = extract_document_b_sections(document_bytes(document))

    assert len(sections) == 1
    assert sections[0].section_id == "bullet-library"
    assert sections[0].section_text == "Content after accidental empty heading."


def test_duplicate_hierarchical_titles_receive_deterministic_suffix() -> None:
    document = Document()
    document.add_heading("Rules", level=1)
    document.add_heading("Purpose", level=2)
    document.add_heading("Purpose", level=2)

    first = extract_document_b_sections(document_bytes(document))
    second = extract_document_b_sections(document_bytes(document))

    assert [section.section_id for section in first] == [
        "rules",
        "rules-purpose",
        "rules-purpose-2",
    ]
    assert first == second


def test_rejects_document_without_recognised_headings() -> None:
    document = Document()
    document.add_paragraph("Ordinary body content.")
    document.add_paragraph("1. Ordinary numbered checklist item.")

    with pytest.raises(DocumentBExtractionError, match="no recognised headings"):
        extract_document_b_sections(document_bytes(document))
