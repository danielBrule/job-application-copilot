"""Shared pytest fixtures."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from job_application_copilot.config.document_b_routing import (
    load_document_b_routing_config,
)
from job_application_copilot.config.settings import AppSettings
from job_application_copilot.ui.dependencies import get_database

DOCUMENT_B_ROUTING_TEMPLATE = (
    Path(__file__).parents[1] / "templates" / "document-b-lane-routes.template.yaml"
)


@pytest.fixture(autouse=True)
def clear_ui_database_cache() -> None:
    """Isolate Streamlit's cached database resource between tests."""

    get_database.clear()
    yield
    get_database.clear()


def make_routable_document_b(
    text: str = "Configured section content.",
    *,
    omitted_heading_paths: frozenset[tuple[str, ...]] = frozenset(),
) -> bytes:
    """Build a complete synthetic DOCX from the canonical exact heading paths."""

    config = load_document_b_routing_config(DOCUMENT_B_ROUTING_TEMPLATE)
    tree: dict[str, dict] = {}
    for catalog in config.section_catalog.values():
        if catalog.heading_path in omitted_heading_paths:
            continue
        node = tree
        for heading in catalog.heading_path:
            node = node.setdefault(heading, {})

    document = Document()

    def add_nodes(nodes: dict[str, dict], level: int) -> None:
        for heading, children in nodes.items():
            document.add_heading(heading, level=level)
            document.add_paragraph(text)
            add_nodes(children, level + 1)

    add_nodes(tree, 1)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def install_document_b_routing_config(settings: AppSettings) -> Path:
    """Install the committed routing template in a test's private data directory."""

    config_path = settings.document_b_routing_config_path
    config_path.parent.mkdir(parents=True, exist_ok=True)
    config_path.write_bytes(DOCUMENT_B_ROUTING_TEMPLATE.read_bytes())
    return config_path
