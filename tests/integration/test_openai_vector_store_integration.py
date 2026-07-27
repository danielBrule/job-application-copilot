"""Opt-in integration test for the real Document B vector-store lifecycle."""

import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from job_application_copilot.config import AppSettings
from job_application_copilot.domain import (
    ReferenceAssetProcessingStatus,
    ReferenceAssetType,
)
from job_application_copilot.llm import OpenAIClient
from job_application_copilot.repositories import create_database
from job_application_copilot.repositories.reference_asset_repository import (
    ReferenceAssetRepository,
)
from job_application_copilot.services import (
    DocumentBVectorStoreService,
    OpenAIFileUploadService,
    ReferenceAssetStorageService,
)
from job_application_copilot.services.database_bootstrap import initialize_database

pytestmark = [
    pytest.mark.openai_integration,
    pytest.mark.skipif(
        os.getenv("JAC_RUN_OPENAI_INTEGRATION") != "1",
        reason="Set JAC_RUN_OPENAI_INTEGRATION=1 to contact OpenAI.",
    ),
]


def make_docx() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_heading("CV generation and positioning guidance", level=1)
    document.add_paragraph(
        "Temporary Document B content for the Job Application Copilot integration test."
    )
    document.save(buffer)
    return buffer.getvalue()


def test_indexes_searches_activates_and_cleans_up_real_vector_store(
    tmp_path: Path,
) -> None:
    data_dir = tmp_path / "data"
    settings = AppSettings(
        _env_file=".env",
        data_dir=data_dir,
        database_path=data_dir / "database" / "integration.db",
        reference_folder=data_dir / "reference",
    )
    settings.database_path.parent.mkdir(parents=True)
    initialize_database(settings.database_path)
    database = create_database(settings.database_path)
    client = OpenAIClient.from_settings(settings)
    remote_file_id: str | None = None
    vector_store_id: str | None = None
    candidate_key: str | None = None
    candidate_version: int | None = None

    try:
        candidate = ReferenceAssetStorageService(database, settings).replace(
            filename="integration-document-b.docx",
            content=make_docx(),
            asset_key="document-b",
            asset_type=ReferenceAssetType.DOCUMENT,
            name="Document B integration test",
        )
        candidate_key = candidate.asset_key
        candidate_version = candidate.version
        uploaded = OpenAIFileUploadService(database, settings, client).upload(
            candidate.asset_key,
            candidate.version,
        )
        remote_file_id = uploaded.openai_file_id

        activated = DocumentBVectorStoreService(database, settings, client).process(
            candidate.asset_key,
            candidate.version,
        )
        vector_store_id = activated.openai_vector_store_id

        assert remote_file_id is not None
        assert vector_store_id is not None
        assert activated.processing_status is ReferenceAssetProcessingStatus.READY
        assert activated.is_active
        assert activated.openai_vector_store_usage_bytes is not None
        assert activated.openai_vector_store_usage_bytes >= 0
    finally:
        if candidate_key is not None and candidate_version is not None:
            with database.session() as session:
                stored = ReferenceAssetRepository(session).require_version(
                    candidate_key,
                    candidate_version,
                )
                remote_file_id = remote_file_id or stored.openai_file_id
                vector_store_id = vector_store_id or stored.openai_vector_store_id
        if vector_store_id is not None:
            client.delete_vector_store(vector_store_id)
        if remote_file_id is not None:
            client.delete_file(remote_file_id)
        client.close()
        database.dispose()
