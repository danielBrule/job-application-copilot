"""Tests for explicit cleanup of inactive tracked OpenAI resources."""

import hashlib
from io import BytesIO
from pathlib import Path
from typing import cast
from unittest.mock import Mock

import pytest
from conftest import make_routable_document_b
from docx import Document

from job_application_copilot.config import AppSettings
from job_application_copilot.domain import (
    ReferenceAssetProcessingStatus,
    ReferenceAssetType,
)
from job_application_copilot.llm import (
    OpenAIClient,
    OpenAIClientError,
    OpenAIConfigurationError,
    OpenAIVectorStore,
    OpenAIVectorStoreFile,
    OpenAIVectorStoreFileStatus,
    OpenAIVectorStoreSearchResult,
    UploadedOpenAIFile,
)
from job_application_copilot.repositories import Database, create_database
from job_application_copilot.repositories.models import ReferenceAsset
from job_application_copilot.repositories.reference_asset_repository import (
    ReferenceAssetRepository,
)
from job_application_copilot.services import (
    ReferenceAssetRemoteCleanupError,
    ReferenceAssetRemoteCleanupNotAllowedError,
    ReferenceAssetRemoteCleanupService,
    ReferenceAssetRemoteRestoreError,
)
from job_application_copilot.services.database_bootstrap import initialize_database
from job_application_copilot.services.remote_reference_operation import (
    remote_reference_operation,
)


@pytest.fixture
def cleanup_context(
    tmp_path: Path,
) -> tuple[ReferenceAssetRemoteCleanupService, Database, AppSettings, Mock]:
    settings = AppSettings(_env_file=None, data_dir=tmp_path / "data")
    settings.database_path.parent.mkdir(parents=True)
    initialize_database(settings.database_path)
    database = create_database(settings.database_path)
    cleaner = Mock(spec=OpenAIClient)
    service = ReferenceAssetRemoteCleanupService(
        database,
        settings,
        client_factory=lambda _: cast(OpenAIClient, cleaner),
    )
    try:
        yield service, database, settings, cleaner
    finally:
        database.dispose()


def add_asset(
    database: Database,
    settings: AppSettings,
    *,
    asset_key: str,
    version: int,
    active: bool = False,
    status: ReferenceAssetProcessingStatus = ReferenceAssetProcessingStatus.READY,
    file_id: str | None = None,
    vector_store_id: str | None = None,
    usage_bytes: int | None = None,
) -> Path:
    relative_path = f"document_b/{asset_key}-v{version:04d}.docx"
    local_path = settings.reference_folder / relative_path
    local_path.parent.mkdir(parents=True, exist_ok=True)
    if asset_key == "document-b":
        content = make_routable_document_b(f"{asset_key}-{version}")
    else:
        document = Document()
        document.add_heading("CV generation and positioning guidance", level=1)
        document.add_paragraph(f"{asset_key}-{version}")
        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
    local_path.write_bytes(content)
    with database.session() as session:
        session.add(
            ReferenceAsset(
                asset_key=asset_key,
                asset_type=ReferenceAssetType.DOCUMENT,
                name="Document A" if asset_key == "document-a" else "Document B",
                version=version,
                file_path=relative_path,
                file_hash=f"sha256:{hashlib.sha256(content).hexdigest()}",
                processing_status=status,
                is_active=active,
                openai_file_id=file_id,
                openai_vector_store_id=vector_store_id,
                openai_vector_store_usage_bytes=usage_bytes,
            )
        )
    return local_path


def test_lists_only_inactive_non_processing_versions_with_remote_resources(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, _ = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
        file_id="file_a_old",
    )
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=2,
        active=True,
        file_id="file_a_active",
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        status=ReferenceAssetProcessingStatus.FAILED,
        file_id="file_b_old",
        vector_store_id="vs_b_old",
        usage_bytes=8_192,
    )
    add_asset(
        database,
        settings,
        asset_key="document-b-processing",
        version=1,
        status=ReferenceAssetProcessingStatus.PROCESSING,
        file_id="file_processing",
    )
    add_asset(
        database,
        settings,
        asset_key="document-local-only",
        version=1,
    )

    candidates = service.list_candidates()

    assert [
        (
            candidate.asset_key,
            candidate.version,
            candidate.openai_file_id,
            candidate.openai_vector_store_id,
        )
        for candidate in candidates
    ] == [
        ("document-a", 1, "file_a_old", None),
        ("document-b", 1, "file_b_old", "vs_b_old"),
    ]
    assert candidates[1].openai_vector_store_usage_bytes == 8_192


def test_deletes_document_b_store_then_file_and_retains_local_history(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, cleaner = cleanup_context
    local_path = add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        file_id="file_old",
        vector_store_id="vs_old",
        usage_bytes=8_192,
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=2,
        active=True,
        file_id="file_active",
        vector_store_id="vs_active",
        usage_bytes=16_384,
    )

    result = service.cleanup("document-b", 1)

    assert result.vector_store_deleted
    assert result.file_deleted
    assert [call[0] for call in cleaner.method_calls] == [
        "delete_vector_store",
        "delete_file",
        "close",
    ]
    with database.session() as session:
        repository = ReferenceAssetRepository(session)
        cleaned = repository.require_version("document-b", 1)
        active = repository.require_version("document-b", 2)
        assert cleaned.openai_vector_store_id is None
        assert cleaned.openai_vector_store_usage_bytes is None
        assert cleaned.openai_file_id is None
        assert cleaned.processing_status is ReferenceAssetProcessingStatus.READY
        assert not cleaned.is_active
        assert active.openai_vector_store_id == "vs_active"
        assert active.openai_file_id == "file_active"
        assert active.is_active
    assert local_path.exists()


def test_deletes_only_document_a_file(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, cleaner = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
        file_id="file_a_old",
    )

    result = service.cleanup("document-a", 1)

    assert not result.vector_store_deleted
    assert result.file_deleted
    cleaner.delete_vector_store.assert_not_called()
    cleaner.delete_file.assert_called_once_with("file_a_old")


def test_file_failure_keeps_file_association_after_store_cleanup(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, cleaner = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        status=ReferenceAssetProcessingStatus.FAILED,
        file_id="file_failed",
        vector_store_id="vs_failed",
        usage_bytes=4_096,
    )
    cleaner.delete_file.side_effect = OpenAIClientError(
        "OpenAI could not be reached after the configured retries.",
        operation="delete",
        retryable=True,
    )

    with pytest.raises(ReferenceAssetRemoteCleanupError, match="could not be reached"):
        service.cleanup("document-b", 1)

    with database.session() as session:
        candidate = ReferenceAssetRepository(session).require_version("document-b", 1)
        assert candidate.openai_vector_store_id is None
        assert candidate.openai_vector_store_usage_bytes is None
        assert candidate.openai_file_id == "file_failed"
    cleaner.close.assert_called_once()


@pytest.mark.parametrize(
    ("active", "status", "expected"),
    [
        (True, ReferenceAssetProcessingStatus.READY, "is active"),
        (False, ReferenceAssetProcessingStatus.PROCESSING, "is processing"),
    ],
)
def test_rejects_unsafe_cleanup_targets(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
    active: bool,
    status: ReferenceAssetProcessingStatus,
    expected: str,
) -> None:
    service, database, settings, cleaner = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        active=active,
        status=status,
        file_id="file_b",
        vector_store_id="vs_b",
    )

    with pytest.raises(ReferenceAssetRemoteCleanupNotAllowedError, match=expected):
        service.cleanup("document-b", 1)

    cleaner.delete_vector_store.assert_not_called()
    cleaner.delete_file.assert_not_called()


def test_rejects_identifier_shared_with_active_asset(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, cleaner = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
        file_id="file_shared",
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        active=True,
        file_id="file_shared",
        vector_store_id="vs_active",
    )

    with pytest.raises(
        ReferenceAssetRemoteCleanupNotAllowedError,
        match="active reference asset",
    ):
        service.cleanup("document-a", 1)

    cleaner.delete_file.assert_not_called()


def test_configuration_failure_releases_operation_guard(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    _, database, settings, _ = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
        file_id="file_a",
    )

    def fail_configuration(_: AppSettings) -> OpenAIClient:
        raise OpenAIConfigurationError("OPENAI_API_KEY is required for OpenAI file operations.")

    service = ReferenceAssetRemoteCleanupService(
        database,
        settings,
        client_factory=fail_configuration,
    )

    with pytest.raises(ReferenceAssetRemoteCleanupError, match="OPENAI_API_KEY"):
        service.cleanup("document-a", 1)

    with remote_reference_operation(
        settings,
        lambda _: cast(OpenAIClient, Mock()),
        RuntimeError,
    ):
        pass


def test_lists_retained_document_versions_without_remote_resources(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, _ = cleanup_context
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
        status=ReferenceAssetProcessingStatus.FAILED,
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=2,
        active=True,
        file_id="file_active",
        vector_store_id="vs_active",
    )

    candidates = service.list_restorable_versions()

    assert [
        (candidate.asset_key, candidate.version, candidate.processing_status)
        for candidate in candidates
    ] == [
        ("document-a", 1, ReferenceAssetProcessingStatus.READY),
        ("document-b", 1, ReferenceAssetProcessingStatus.FAILED),
    ]


def test_restores_and_activates_retained_document_a(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, client = cleanup_context
    local_path = add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
    )
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=2,
        active=True,
        file_id="file_current",
    )
    client.upload_docx.return_value = UploadedOpenAIFile(
        file_id="file_restored",
        filename="document-a-v0001.docx",
        size_bytes=local_path.stat().st_size,
        request_id="req_restore",
    )

    restored = service.restore("document-a", 1)

    assert restored.is_active
    assert restored.processing_status is ReferenceAssetProcessingStatus.READY
    assert restored.openai_file_id == "file_restored"
    client.create_vector_store.assert_not_called()
    with database.session() as session:
        repository = ReferenceAssetRepository(session)
        previous = repository.require_version("document-a", 2)
        assert not previous.is_active
        assert previous.openai_file_id == "file_current"
    assert local_path.exists()


def test_restores_indexes_and_activates_retained_document_b(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, client = cleanup_context
    local_path = add_asset(
        database,
        settings,
        asset_key="document-b",
        version=1,
    )
    add_asset(
        database,
        settings,
        asset_key="document-b",
        version=2,
        active=True,
        file_id="file_current",
        vector_store_id="vs_current",
    )
    client.upload_docx.return_value = UploadedOpenAIFile(
        file_id="file_restored",
        filename="document-b-v0001.docx",
        size_bytes=local_path.stat().st_size,
        request_id="req_upload",
    )
    client.create_vector_store.return_value = OpenAIVectorStore(
        vector_store_id="vs_restored",
        status="in_progress",
        usage_bytes=0,
        request_id="req_store",
    )
    client.wait_for_vector_store_file.return_value = OpenAIVectorStoreFile(
        file_id="file_restored",
        vector_store_id="vs_restored",
        status=OpenAIVectorStoreFileStatus.COMPLETED,
        usage_bytes=4_096,
        error_code=None,
        request_id="req_poll",
    )
    client.search_vector_store.return_value = (
        OpenAIVectorStoreSearchResult(
            file_id="file_restored",
            filename="document-b-v0001.docx",
            score=0.9,
            text="CV generation and positioning guidance.",
        ),
    )

    restored = service.restore("document-b", 1)

    assert restored.is_active
    assert restored.openai_file_id == "file_restored"
    assert restored.openai_vector_store_id == "vs_restored"
    assert restored.openai_vector_store_usage_bytes == 4_096
    with database.session() as session:
        previous = ReferenceAssetRepository(session).require_version("document-b", 2)
        assert not previous.is_active
        assert previous.openai_vector_store_id == "vs_current"


def test_restoration_integrity_failure_preserves_current_active_version(
    cleanup_context: tuple[
        ReferenceAssetRemoteCleanupService,
        Database,
        AppSettings,
        Mock,
    ],
) -> None:
    service, database, settings, client = cleanup_context
    retained_path = add_asset(
        database,
        settings,
        asset_key="document-a",
        version=1,
    )
    add_asset(
        database,
        settings,
        asset_key="document-a",
        version=2,
        active=True,
        file_id="file_current",
    )
    retained_path.write_bytes(b"changed after storage")

    with pytest.raises(ReferenceAssetRemoteRestoreError, match="recorded hash"):
        service.restore("document-a", 1)

    client.upload_docx.assert_not_called()
    with database.session() as session:
        current = ReferenceAssetRepository(session).require_version("document-a", 2)
        assert current.is_active
        assert current.openai_file_id == "file_current"
