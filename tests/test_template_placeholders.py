"""Tests for English template placeholder discovery and manifest validation."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError

from job_application_copilot.config import AppSettings
from job_application_copilot.documents.template_placeholders import extract_template_placeholders
from job_application_copilot.domain import (
    ENGLISH_CV_TEMPLATE_KEY,
    CvTemplateManifest,
    CvTemplateManifestStatus,
    CvTemplateSlotKind,
    CvTemplateSlotMapping,
    ReferenceAssetType,
)
from job_application_copilot.repositories import Database, create_database
from job_application_copilot.repositories.reference_asset_repository import (
    ReferenceAssetRepository,
)
from job_application_copilot.services import CvTemplateManifestService, ReferenceAssetStorageService
from job_application_copilot.services.database_bootstrap import initialize_database
from job_application_copilot.ui.components.template_manifest_settings import (
    _experience_target,
    _suggest_mapping,
)


def template_bytes(suffix: str = "") -> bytes:
    document = Document()
    document.add_paragraph("[OPENING_TITLE]")
    if suffix:
        document.add_paragraph(suffix)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "[EXPERIENCE_CURRENT]"
    table.cell(0, 1).text = "[EXPERIENCE_CURRENT] [SKILLS]"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_discovers_unique_placeholders_in_paragraphs_and_tables() -> None:
    assert extract_template_placeholders(template_bytes()) == (
        "[OPENING_TITLE]",
        "[EXPERIENCE_CURRENT]",
        "[SKILLS]",
    )


def test_confirmed_manifest_requires_every_discovered_placeholder() -> None:
    with pytest.raises(ValidationError, match="map every discovered placeholder"):
        CvTemplateManifest(
            template_asset_id=1,
            status=CvTemplateManifestStatus.CONFIRMED,
            placeholders=("[OPENING_TITLE]", "[SKILLS]"),
            slots=(
                CvTemplateSlotMapping(
                    placeholder="[OPENING_TITLE]", kind=CvTemplateSlotKind.OPENING_TITLE
                ),
            ),
        )


def test_experience_mapping_requires_a_target() -> None:
    with pytest.raises(ValidationError, match="experience target"):
        CvTemplateSlotMapping(
            placeholder="[EXPERIENCE_CURRENT]", kind=CvTemplateSlotKind.EXPERIENCE
        )


@pytest.mark.parametrize(
    ("placeholder", "kind", "target"),
    [
        ("[OPENING_TITLE]", CvTemplateSlotKind.OPENING_TITLE, ""),
        ("[OPENING_PROFILE]", CvTemplateSlotKind.OPENING_PROFILE, ""),
        ("[SKILLS]", CvTemplateSlotKind.SKILLS, ""),
        ("[EXPERIENCE_EKIMETRICS]", CvTemplateSlotKind.EXPERIENCE, "Ekimetrics"),
        ("[EKIMETRICS_TITLE]", CvTemplateSlotKind.EXPERIENCE_TITLE, "Ekimetrics"),
    ],
)
def test_suggests_editable_template_mapping_defaults(
    placeholder: str, kind: CvTemplateSlotKind, target: str
) -> None:
    assert _suggest_mapping(placeholder) == (kind, target)


def test_derives_experience_target_when_the_user_changes_a_placeholder_type() -> None:
    assert _experience_target("[ROLE_HISTORY]", "") == "Role History"


@pytest.fixture
def manifest_service(tmp_path: Path) -> tuple[CvTemplateManifestService, Database]:
    settings = AppSettings(_env_file=None, data_dir=tmp_path / "data")
    settings.database_path.parent.mkdir(parents=True)
    initialize_database(settings.database_path)
    database = create_database(settings.database_path)
    try:
        yield CvTemplateManifestService(database, settings), database
    finally:
        database.dispose()


def test_candidate_template_remains_inactive_until_valid_mapping_is_confirmed(
    manifest_service: tuple[CvTemplateManifestService, Database],
) -> None:
    service, database = manifest_service
    existing = ReferenceAssetStorageService(database, service.settings).replace(
        filename="existing.docx",
        content=template_bytes(),
        asset_key=ENGLISH_CV_TEMPLATE_KEY,
        asset_type=ReferenceAssetType.TEMPLATE,
        name="English CV template",
        language_code="en",
    )

    candidate = service.upload(filename="candidate.docx", content=template_bytes("candidate"))

    with database.session() as session:
        versions = ReferenceAssetRepository(session).list_versions(ENGLISH_CV_TEMPLATE_KEY)
        assert [(version.id, version.is_active) for version in versions] == [
            (candidate.template_asset_id, False),
            (existing.id, True),
        ]

    confirmed = service.confirm(
        template_asset_id=candidate.template_asset_id,
        slots=(
            CvTemplateSlotMapping(
                placeholder="[OPENING_TITLE]", kind=CvTemplateSlotKind.OPENING_TITLE
            ),
            CvTemplateSlotMapping(
                placeholder="[EXPERIENCE_CURRENT]",
                kind=CvTemplateSlotKind.EXPERIENCE,
                experience_target="Current employer",
            ),
            CvTemplateSlotMapping(placeholder="[SKILLS]", kind=CvTemplateSlotKind.SKILLS),
        ),
    )

    assert confirmed.status is CvTemplateManifestStatus.CONFIRMED
    with database.session() as session:
        versions = ReferenceAssetRepository(session).list_versions(ENGLISH_CV_TEMPLATE_KEY)
        assert [(version.id, version.is_active) for version in versions] == [
            (candidate.template_asset_id, True),
            (existing.id, False),
        ]
