"""Discover bracketed generated-content placeholders in DOCX templates."""

import re
from io import BytesIO

from docx import Document

PLACEHOLDER_FIND_PATTERN = re.compile(r"\[[A-Z][A-Z0-9_]*\]")


def extract_template_placeholders(content: bytes) -> tuple[str, ...]:
    """Return unique body paragraph and table placeholders in first-seen order."""

    document = Document(BytesIO(content))
    found: list[str] = []
    for paragraph in document.paragraphs:
        found.extend(_placeholders(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                found.extend(_placeholders(cell.text))
    return tuple(dict.fromkeys(found))


def _placeholders(text: str) -> tuple[str, ...]:
    return tuple(match.group(0) for match in PLACEHOLDER_FIND_PATTERN.finditer(text))
