"""Deterministically populate generated content into a DOCX CV template."""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from job_application_copilot.documents.docx_validation import validate_docx
from job_application_copilot.domain import CvTemplateManifest, CvTemplateSlotKind, FinalCvOutput
from job_application_copilot.errors import ApplicationValidationError


class CvDocumentRenderError(ApplicationValidationError):
    """Raised when generated content cannot safely populate a DOCX template."""


def render_cv_template(
    template_content: bytes,
    *,
    manifest: CvTemplateManifest,
    output: FinalCvOutput,
) -> bytes:
    """Return a populated DOCX copy without changing template-owned content or formatting."""

    slot_content = _slot_content(manifest, output)
    document = Document(BytesIO(template_content))
    populated: set[str] = set()

    for paragraph in _paragraphs(document):
        matching = tuple(
            placeholder for placeholder in slot_content if placeholder in paragraph.text
        )
        if not matching:
            continue
        if len(matching) != 1:
            raise CvDocumentRenderError(
                "A generated template paragraph must contain exactly one configured placeholder."
            )
        placeholder = matching[0]
        if placeholder != paragraph.text:
            raise CvDocumentRenderError(
                f"Template placeholder {placeholder} must occupy its paragraph without surrounding text."
            )
        _replace_paragraph(paragraph, slot_content[placeholder])
        populated.add(placeholder)

    missing = set(slot_content).difference(populated)
    if missing:
        names = ", ".join(sorted(missing))
        raise CvDocumentRenderError(
            f"Template is missing configured generated placeholders: {names}."
        )

    unresolved = _unresolved_placeholders(document, set(manifest.placeholders))
    if unresolved:
        names = ", ".join(sorted(unresolved))
        raise CvDocumentRenderError(f"Rendered CV contains unresolved placeholders: {names}.")

    buffer = BytesIO()
    document.save(buffer)
    rendered_content = buffer.getvalue()
    validate_docx("rendered-cv.docx", rendered_content)
    return rendered_content


def _slot_content(
    manifest: CvTemplateManifest, output: FinalCvOutput
) -> dict[str, tuple[str, ...]]:
    values: dict[str, tuple[str, ...]] = {
        output.opening_title.placeholder: (output.opening_title.content,),
        output.opening_profile.placeholder: (output.opening_profile.content,),
        output.skills.placeholder: tuple(
            f"{entry.name}: {entry.content}" for entry in output.skills.entries
        ),
    }
    experience_by_placeholder = {item.placeholder: item for item in output.experience}
    titles_by_target = {
        slot.experience_target: slot.placeholder
        for slot in manifest.slots
        if slot.kind is CvTemplateSlotKind.EXPERIENCE_TITLE
    }
    for slot in manifest.slots:
        if slot.kind is not CvTemplateSlotKind.EXPERIENCE:
            continue
        experience = experience_by_placeholder[slot.placeholder]
        values[slot.placeholder] = tuple(
            value for value in (experience.introduction, *experience.bullets) if value is not None
        )
        title_placeholder = titles_by_target.get(slot.experience_target)
        if title_placeholder is not None:
            values[title_placeholder] = (
                () if experience.title is None else (experience.title.content,)
            )
    return values


def _paragraphs(document: DocumentType) -> tuple[Paragraph, ...]:
    paragraphs: list[Paragraph] = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(_table_paragraphs(table))
    return tuple(paragraphs)


def _table_paragraphs(table: Table) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(_cell_paragraphs(cell))
    return paragraphs


def _cell_paragraphs(cell: _Cell) -> list[Paragraph]:
    paragraphs = list(cell.paragraphs)
    for table in cell.tables:
        paragraphs.extend(_table_paragraphs(table))
    return paragraphs


def _replace_paragraph(paragraph: Paragraph, values: tuple[str, ...]) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise CvDocumentRenderError("Template paragraph has no XML parent.")
    insertion_point: Any = paragraph._p
    for value in values:
        replacement = _cloned_paragraph(paragraph, value)
        insertion_point.addnext(replacement)
        insertion_point = replacement
    parent.remove(paragraph._p)


def _cloned_paragraph(paragraph: Paragraph, value: str) -> object:
    replacement = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        replacement.append(deepcopy(paragraph._p.pPr))
    run = OxmlElement("w:r")
    if paragraph._p.r_lst and paragraph._p.r_lst[0].rPr is not None:
        run.append(deepcopy(paragraph._p.r_lst[0].rPr))
    text = OxmlElement("w:t")
    if value[:1].isspace() or value[-1:].isspace():
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value
    run.append(text)
    replacement.append(run)
    return replacement


def _unresolved_placeholders(document: DocumentType, placeholders: set[str]) -> set[str]:
    return {
        placeholder
        for paragraph in _paragraphs(document)
        for placeholder in placeholders
        if placeholder in paragraph.text
    }
