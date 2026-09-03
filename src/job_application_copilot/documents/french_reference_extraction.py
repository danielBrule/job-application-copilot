"""Extract searchable paragraph and table text from a French reference DOCX."""

from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import OpcError
from lxml.etree import XMLSyntaxError


class FrenchReferenceExtractionError(ValueError):
    """Raised when a validated DOCX has no searchable text."""


def extract_french_reference_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except (BadZipFile, KeyError, OpcError, ValueError, XMLSyntaxError) as error:
        raise FrenchReferenceExtractionError(
            "The French reference is not a readable DOCX."
        ) from error
    blocks = [" ".join(paragraph.text.split()) for paragraph in document.paragraphs]
    blocks.extend(
        " | ".join(" ".join(cell.text.split()) for cell in row.cells)
        for table in document.tables
        for row in table.rows
    )
    text = "\n\n".join(block for block in blocks if block)
    if not text:
        raise FrenchReferenceExtractionError("The French reference contains no searchable text.")
    return text
